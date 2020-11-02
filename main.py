import re
import pythaispell
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

class Service:
    def start(self):
        print('[*] Start spell checking...')
        raw = self.readDoc()
        result = self.spellChecking(raw)
        arrayString = self.splitText(result)
        self.writeDoc(arrayString)
        print('[*] Spell checking done...')

    def readDoc(self):
        doc = Document('input/Document.docx')
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    def writeDoc(self, textList):
        print('[*] Start writing file in output floder...')
        doc = Document()
        p = doc.add_paragraph("")
        for i, e in enumerate(textList):
            if '<คำผิด>' in e:
                word = e.replace('<คำผิด>', "")
                word = word.replace('</คำผิด>', "")
                font = p.add_run(word).font.highlight_color = WD_COLOR_INDEX.YELLOW
                continue
            p.add_run(e)
        doc.save('output/Document.docx')
        print('[*] Success on writing file in output floder...')

    def spellChecking(self, text):
        print('[*] Start spell checking...')
        result = pythaispell.spell(text)
        print('[*] Success on spell checking...')
        return result

    def splitText(self, text):
        return re.split('(<คำผิด>.*<\/คำผิด>)', text)
    
service = Service()
service.start()